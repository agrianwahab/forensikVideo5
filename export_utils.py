# export_utils.py
# VERSI BARU - FOKUS PADA STABILITAS PDF DAN DOCX

import io
from pathlib import Path
from typing import Optional, Any
import streamlit as st
from datetime import datetime

def check_dependency(package_name: str) -> bool:
    """Memeriksa apakah sebuah library Python terpasang."""
    import importlib.util
    return importlib.util.find_spec(package_name) is not None

def create_docx_report_robust(result: Any, output_path: Path) -> Optional[Path]:
    """
    Membuat laporan DOCX ringkas yang sangat aman dan anti-gagal.
    Hanya memasukkan data yang ada untuk menghindari error.
    """
    if not check_dependency('docx'):
        print("Peringatan: `python-docx` tidak terpasang. Ekspor DOCX dilewati.")
        return None

    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        style = doc.styles['Normal'].font
        style.name = 'Arial'
        style.size = Pt(11)

        # Judul Laporan
        doc.add_heading('Laporan Ringkas Forensik Video VIFA-Pro', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Dihasilkan pada: {datetime.now().strftime('%d %B %Y, %H:%M:%S')}", style='Caption').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        # --- Bagian 1: Informasi Kasus (Sangat Defensif) ---
        doc.add_heading('1. Informasi Bukti', level=1)
        video_name = Path(getattr(result, 'video_path', 'N/A')).name
        preservation_hash = getattr(result, 'preservation_hash', 'N/A')
        summary = getattr(result, 'summary', {})
        total_frames = summary.get('total_frames', 'N/A')
        total_anomaly = summary.get('total_anomaly', 'N/A')

        info_data = {
            "Nama File Bukti": video_name,
            "Hash Integritas (SHA-256)": preservation_hash,
            "Total Frame Dianalisis": str(total_frames),
            "Total Anomali Ditemukan": str(total_anomaly),
        }
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.cell(0, 0).text = "Item"
        table.cell(0, 1).text = "Detail"
        for key, value in info_data.items():
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = value
        
        doc.add_paragraph()

        # --- Bagian 2: Hasil Analisis FERM (Sangat Defensif) ---
        doc.add_heading('2. Penilaian Keandalan Bukti (FERM)', level=1)
        ferm = getattr(result, 'forensic_evidence_matrix', {})
        conclusion = ferm.get('conclusion', {})
        
        reliability = conclusion.get('reliability_assessment', 'Tidak Dapat Ditentukan')
        p = doc.add_paragraph()
        p.add_run('Penilaian Keseluruhan: ').bold = True
        p.add_run(reliability)
        
        findings = conclusion.get('primary_findings', [])
        if findings:
            doc.add_paragraph('Temuan Kunci:', style='Intense Quote')
            for item in findings:
                finding_text = item.get('finding', 'N/A')
                confidence = item.get('confidence', 'N/A')
                doc.add_paragraph(f"• {finding_text} (Kepercayaan: {confidence})", style='List Bullet')
        
        recommendations = conclusion.get('recommended_actions', [])
        if recommendations:
            doc.add_paragraph('Rekomendasi Tindak Lanjut:', style='Intense Quote')
            for action in recommendations:
                doc.add_paragraph(f"• {action}", style='List Bullet')
        
        doc.add_page_break()
        
        # --- Bagian 3: Detail Peristiwa Anomali ---
        doc.add_heading('3. Detail Peristiwa Anomali', level=1)
        localizations = getattr(result, 'localizations', [])
        if not localizations:
            doc.add_paragraph("Tidak ada peristiwa anomali signifikan yang ditemukan.")
        else:
            for i, loc in enumerate(localizations):
                event_type = loc.get('event', 'unknown').replace('anomaly_', '').capitalize()
                start_ts = loc.get('start_ts', 0)
                end_ts = loc.get('end_ts', 0)
                confidence = loc.get('confidence', 'N/A')
                doc.add_heading(f"Peristiwa #{i+1}: {event_type} @ {start_ts:.2f}s", level=2)
                doc.add_paragraph(f"Durasi: {end_ts - start_ts:.2f} detik | Kepercayaan: {confidence}")
                
                if isinstance(loc.get('metrics'), dict):
                    doc.add_paragraph("Metrik Terdeteksi:")
                    for key, val in loc['metrics'].items():
                         doc.add_paragraph(f"  - {key.replace('_', ' ').title()}: {val}", style='List Paragraph')
        
        doc.save(output_path)
        return output_path
    except Exception as e:
        print(f"FATAL: Terjadi error saat membuat file DOCX: {e}")
        import traceback
        traceback.print_exc()
        return None


def add_export_buttons(pdf_path: Path, result: Any, col1, col2, col3):
    """
    Menampilkan tombol ekspor dengan UI yang lebih sederhana dan fokus.
    Kolom PNG dihilangkan untuk saat ini.
    """
    # Kolom 1: Tombol Download PDF (Wajib)
    with col1:
        if pdf_path.exists():
            with open(pdf_path, "rb") as f:
                st.download_button(
                    label="📄 Download PDF",
                    data=f.read(),
                    file_name=pdf_path.name,
                    mime="application/pdf",
                    use_container_width=True
                )
        else:
            st.error("Laporan PDF tidak berhasil dibuat. Periksa log konsol.")

    # Kolom 2: Tombol Download DOCX
    with col2:
        # Cari file docx di dalam atribut result
        docx_path_str = getattr(result, 'docx_report_path', None)
        if docx_path_str and Path(docx_path_str).exists():
             with open(docx_path_str, "rb") as f:
                st.download_button(
                    label="📝 Download DOCX",
                    data=f.read(),
                    file_name=Path(docx_path_str).name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
        else:
            st.warning("DOCX tidak dibuat. Cek log konsol untuk detailnya.", icon="⚠️")

    # Kolom 3 Dibiarkan kosong atau untuk penggunaan lain.
    with col3:
        st.write("") # Placeholder